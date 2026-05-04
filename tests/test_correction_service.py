"""Tests for correction_service module."""

import pytest
from pathlib import Path
from unittest.mock import Mock, patch, MagicMock
from uuid import uuid4
import tempfile
import shutil

from app.domain.services.correction_service import CorrectionService
from app.domain.entities.validation_report import (
    ValidationReport, ViolationDetail, TextLocation,
    ViolationCategory, ViolationSeverity
)


class TestCorrectionService:
    """Test suite for CorrectionService class."""

    @pytest.fixture
    def temp_dir(self):
        """Create a temporary directory for test files."""
        temp_path = Path(tempfile.mkdtemp())
        yield temp_path
        shutil.rmtree(temp_path, ignore_errors=True)

    @pytest.fixture
    def sample_docx(self, temp_dir):
        """Create a sample .docx file for testing."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not available")

        doc_path = temp_dir / "sample.docx"
        doc = Document()
        doc.add_paragraph("First paragraph text.")
        doc.add_paragraph("Second paragraph with some content.")
        doc.save(str(doc_path))
        return doc_path

    @pytest.fixture
    def validation_report_with_violations(self):
        """Create a validation report with sample violations."""
        report = ValidationReport(
            document_name="test.docx",
            template_name="test_template",
        )

        violation1 = ViolationDetail(
            id=uuid4(),
            category=ViolationCategory.FONT,
            severity=ViolationSeverity.ERROR,
            description="Font should be Songti",
            location=TextLocation(paragraph_index=1, text="First paragraph text."),
            original_content="First paragraph text.",
            suggested_fix="Modified first paragraph text.",
            context_before=None,
            context_after=None,
        )

        violation2 = ViolationDetail(
            id=uuid4(),
            category=ViolationCategory.FONT_SIZE,
            severity=ViolationSeverity.WARNING,
            description="Font size should be 12pt",
            location=TextLocation(paragraph_index=2, text="Second paragraph with some content."),
            original_content="Second paragraph with some content.",
            suggested_fix="Second paragraph modified.",
            context_before="First paragraph text.",
            context_after=None,
        )

        report.add_violation(violation1)
        report.add_violation(violation2)

        return report

    def test_initialization(self):
        """Test CorrectionService initializes correctly."""
        service = CorrectionService()
        assert service._merger is not None

    def test_initialization_with_custom_merger(self):
        """Test CorrectionService initializes with custom merger."""
        mock_merger = Mock()
        service = CorrectionService(merger=mock_merger)
        assert service._merger is mock_merger

    def test_initialization_with_output_dir(self, temp_dir):
        """Test CorrectionService initializes with output directory."""
        service = CorrectionService(output_dir=temp_dir)
        assert service._output_dir == temp_dir

    def test_apply_corrections_raises_on_missing_original(self, temp_dir):
        """Test apply_corrections raises FileNotFoundError for missing original."""
        service = CorrectionService()
        report = ValidationReport(document_name="test.docx")
        report.add_violation(ViolationDetail(
            id=uuid4(),
            category=ViolationCategory.FONT,
            severity=ViolationSeverity.ERROR,
            description="Test",
            location=TextLocation(paragraph_index=1, text="test"),
            original_content="old",
            suggested_fix="new",
        ))

        missing_path = temp_dir / "nonexistent.docx"

        with pytest.raises(FileNotFoundError):
            service.apply_corrections(report, {}, missing_path)

    def test_apply_corrections_raises_on_empty_violations(self, temp_dir, sample_docx):
        """Test apply_corrections raises ValueError for empty violations."""
        service = CorrectionService()
        report = ValidationReport(document_name="test.docx", violations=[])

        with pytest.raises(ValueError, match="no violations"):
            service.apply_corrections(report, {}, sample_docx)

    def test_apply_corrections_basic(self, temp_dir, sample_docx, validation_report_with_violations):
        """Test basic correction application."""
        service = CorrectionService(output_dir=temp_dir)

        result = service.apply_corrections(
            validation_report_with_violations,
            user_modifications={},
            original_doc_path=sample_docx,
        )

        assert result.exists()
        assert result.parent == temp_dir

        # Verify content was modified
        from docx import Document
        doc = Document(str(result))
        texts = [p.text for p in doc.paragraphs]

        assert "Modified first paragraph text." in texts
        assert "Second paragraph modified." in texts

    def test_apply_corrections_with_user_modifications(self, temp_dir, sample_docx, validation_report_with_violations):
        """Test correction with user-modified fixes."""
        service = CorrectionService(output_dir=temp_dir)

        # Override the first violation's fix
        user_mods = {
            validation_report_with_violations.violations[0].id: "User custom fix."
        }

        result = service.apply_corrections(
            validation_report_with_violations,
            user_modifications=user_mods,
            original_doc_path=sample_docx,
        )

        from docx import Document
        doc = Document(str(result))
        texts = [p.text for p in doc.paragraphs]

        # First should use user modification
        assert "User custom fix." in texts
        # Second should use suggested_fix (not user modified)
        assert "Second paragraph modified." in texts

    def test_apply_corrections_skips_empty_original(self, temp_dir, sample_docx):
        """Test that violations with empty original_content are skipped."""
        service = CorrectionService(output_dir=temp_dir)

        report = ValidationReport(document_name="test.docx")

        # Add violation with empty original
        violation = ViolationDetail(
            id=uuid4(),
            category=ViolationCategory.FONT,
            severity=ViolationSeverity.ERROR,
            description="Test",
            location=TextLocation(paragraph_index=1, text="test"),
            original_content="",  # Empty
            suggested_fix="new",
        )
        report.add_violation(violation)

        # Add valid violation
        valid_violation = ViolationDetail(
            id=uuid4(),
            category=ViolationCategory.FONT,
            severity=ViolationSeverity.ERROR,
            description="Test",
            location=TextLocation(paragraph_index=1, text="First paragraph text."),
            original_content="First paragraph text.",
            suggested_fix="Modified.",
        )
        report.add_violation(valid_violation)

        result = service.apply_corrections(report, {}, sample_docx)

        from docx import Document
        doc = Document(str(result))
        texts = [p.text for p in doc.paragraphs]

        # Only the valid violation should be applied
        assert "Modified." in texts

    def test_apply_corrections_skips_empty_suggested_fix(self, temp_dir, sample_docx):
        """Test that violations with empty suggested_fix are skipped."""
        service = CorrectionService(output_dir=temp_dir)

        report = ValidationReport(document_name="test.docx")

        violation = ViolationDetail(
            id=uuid4(),
            category=ViolationCategory.FONT,
            severity=ViolationSeverity.ERROR,
            description="Test",
            location=TextLocation(paragraph_index=1, text="test"),
            original_content="Some text",
            suggested_fix="",  # Empty
        )
        report.add_violation(violation)

        result = service.apply_corrections(report, {}, sample_docx)

        # The empty suggestion should not cause crashes
        assert result.exists()

    def test_batch_corrections_success(self, temp_dir):
        """Test batch corrections with multiple documents."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not available")

        # Create two documents
        doc1_path = temp_dir / "doc1.docx"
        doc2_path = temp_dir / "doc2.docx"

        doc1 = Document()
        doc1.add_paragraph("Document 1 content.")
        doc1.save(str(doc1_path))

        doc2 = Document()
        doc2.add_paragraph("Document 2 content.")
        doc2.save(str(doc2_path))

        # Create reports
        report1 = ValidationReport(document_name="doc1.docx")
        report1.add_violation(ViolationDetail(
            id=uuid4(),
            category=ViolationCategory.FONT,
            severity=ViolationSeverity.ERROR,
            description="Test",
            location=TextLocation(paragraph_index=1, text="Document 1 content."),
            original_content="Document 1 content.",
            suggested_fix="Doc 1 modified.",
        ))

        report2 = ValidationReport(document_name="doc2.docx")
        report2.add_violation(ViolationDetail(
            id=uuid4(),
            category=ViolationCategory.FONT,
            severity=ViolationSeverity.ERROR,
            description="Test",
            location=TextLocation(paragraph_index=1, text="Document 2 content."),
            original_content="Document 2 content.",
            suggested_fix="Doc 2 modified.",
        ))

        service = CorrectionService(output_dir=temp_dir)

        results = service.apply_corrections_batch(
            validation_reports=[report1, report2],
            user_modifications_list=[{}, {}],
            original_doc_paths=[doc1_path, doc2_path],
        )

        assert len(results) == 2
        assert all(r.exists() for r in results)

        # Verify content
        doc1_result = Document(str(results[0]))
        assert "Doc 1 modified." in [p.text for p in doc1_result.paragraphs]

        doc2_result = Document(str(results[1]))
        assert "Doc 2 modified." in [p.text for p in doc2_result.paragraphs]

    def test_batch_corrections_raises_on_mismatched_lengths(self, temp_dir):
        """Test batch corrections raises on mismatched list lengths."""
        service = CorrectionService()

        report = ValidationReport(document_name="test.docx")
        report.add_violation(ViolationDetail(
            id=uuid4(),
            category=ViolationCategory.FONT,
            severity=ViolationSeverity.ERROR,
            description="Test",
            location=TextLocation(paragraph_index=1, text="test"),
            original_content="old",
            suggested_fix="new",
        ))

        with pytest.raises(ValueError, match="same length"):
            service.apply_corrections_batch(
                validation_reports=[report],
                user_modifications_list=[],
                original_doc_paths=[],
            )

    def test_using_ai_word_skill_property(self):
        """Test using_ai_word_skill property delegates to merger."""
        mock_merger = Mock()
        mock_merger.using_ai_word_skill = True
        service = CorrectionService(merger=mock_merger)
        assert service.using_ai_word_skill is True

        mock_merger.using_ai_word_skill = False
        assert service.using_ai_word_skill is False

    def test_build_corrections_from_report(self, validation_report_with_violations):
        """Test _build_corrections extracts correct data."""
        service = CorrectionService()

        corrections = service._build_corrections(
            validation_report_with_violations,
            user_modifications={},
        )

        assert len(corrections) == 2

        # First correction
        assert corrections[0]["original"] == "First paragraph text."
        assert corrections[0]["fixed"] == "Modified first paragraph text."
        assert corrections[0]["paragraph_index"] == 1

        # Second correction
        assert corrections[1]["original"] == "Second paragraph with some content."
        assert corrections[1]["fixed"] == "Second paragraph modified."

    def test_build_corrections_with_user_modifications(self, validation_report_with_violations):
        """Test _build_corrections respects user modifications."""
        service = CorrectionService()

        violation_id = validation_report_with_violations.violations[0].id
        user_mods = {violation_id: "User override."}

        corrections = service._build_corrections(
            validation_report_with_violations,
            user_modifications=user_mods,
        )

        # First should be user override
        assert corrections[0]["fixed"] == "User override."
        # Second should be suggested_fix
        assert corrections[1]["fixed"] == "Second paragraph modified."


if __name__ == "__main__":
    pytest.main([__file__, "-v"])