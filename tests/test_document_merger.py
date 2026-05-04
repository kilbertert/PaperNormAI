"""Tests for document_merger module."""

import pytest
from pathlib import Path
from unittest.mock import Mock, patch, MagicMock
import tempfile
import shutil

from app.infrastructure.docx.document_merger import DocumentMerger


class TestDocumentMerger:
    """Test suite for DocumentMerger class."""

    @pytest.fixture
    def temp_dir(self):
        """Create a temporary directory for test files."""
        temp_path = Path(tempfile.mkdtemp())
        yield temp_path
        shutil.rmtree(temp_path, ignore_errors=True)

    @pytest.fixture
    def sample_docx(self, temp_dir):
        """Create a sample .docx file for testing."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not available")

        doc_path = temp_dir / "sample.docx"
        doc = Document()
        doc.add_paragraph("First paragraph text.")
        doc.add_paragraph("Second paragraph with some content.")
        doc.add_paragraph("Third paragraph here.")
        doc.save(str(doc_path))
        return doc_path

    def test_initialization(self):
        """Test DocumentMerger initializes correctly."""
        merger = DocumentMerger()
        assert merger._output_dir is None
        assert merger._ai_word_skill_available is False

    def test_initialization_with_output_dir(self, temp_dir):
        """Test DocumentMerger initializes with custom output directory."""
        merger = DocumentMerger(output_dir=temp_dir)
        assert merger._output_dir == temp_dir

    def test_merge_raises_on_missing_original(self, temp_dir):
        """Test merge raises FileNotFoundError for missing original."""
        merger = DocumentMerger(output_dir=temp_dir)
        missing_path = temp_dir / "nonexistent.docx"
        corrections = [{"original": "foo", "fixed": "bar"}]

        with pytest.raises(FileNotFoundError):
            merger.merge(missing_path, corrections)

    def test_merge_raises_on_missing_docx_library(self):
        """Test merge raises RuntimeError if python-docx unavailable."""
        with patch("app.infrastructure.docx.document_merger.DOCX_AVAILABLE", False):
            merger = DocumentMerger()
            # Create a fake path that exists
            with patch("pathlib.Path.exists", return_value=True):
                with pytest.raises(RuntimeError, match="python-docx"):
                    merger.merge(Path("test.docx"), [])

    def test_output_path_generation(self, temp_dir, sample_docx):
        """Test output path is generated correctly."""
        merger = DocumentMerger(output_dir=temp_dir)
        corrections = [{"original": "old", "fixed": "new"}]

        result = merger.merge(sample_docx, corrections)
        assert result.parent == temp_dir
        assert result.stem == "sample_corrected"
        assert result.suffix == ".docx"

    def test_output_path_with_existing_file(self, temp_dir, sample_docx):
        """Test output path handles existing corrected file."""
        merger = DocumentMerger(output_dir=temp_dir)

        # Create first corrected file
        first_output = temp_dir / "sample_corrected.docx"
        first_output.write_text("dummy")

        corrections = [{"original": "old", "fixed": "new"}]
        result = merger.merge(sample_docx, corrections)

        assert result.name == "sample_corrected_1.docx"

    def test_basic_replacement_simple(self, temp_dir, sample_docx):
        """Test basic text replacement in paragraph."""
        merger = DocumentMerger(output_dir=temp_dir)
        corrections = [
            {
                "original": "First paragraph text.",
                "fixed": "Modified first paragraph.",
                "paragraph_index": 1,
            }
        ]

        result = merger.merge(sample_docx, corrections)
        assert result.exists()

        # Verify the content was changed
        from docx import Document
        doc = Document(str(result))
        texts = [p.text for p in doc.paragraphs]
        assert "Modified first paragraph." in texts

    def test_basic_replacement_multiple(self, temp_dir, sample_docx):
        """Test multiple corrections are applied."""
        merger = DocumentMerger(output_dir=temp_dir)
        corrections = [
            {"original": "First paragraph text.", "fixed": "First modified."},
            {"original": "Second paragraph with some content.", "fixed": "Second modified."},
        ]

        result = merger.merge(sample_docx, corrections)

        from docx import Document
        doc = Document(str(result))
        texts = [p.text for p in doc.paragraphs]
        assert "First modified." in texts
        assert "Second modified." in texts

    def test_paragraph_index_hints_used(self, temp_dir):
        """Test that paragraph_index is used for targeting."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not available")

        doc_path = temp_dir / "indexed.docx"
        doc = Document()
        doc.add_paragraph("Para 1 content.")
        doc.add_paragraph("Para 2 content.")
        doc.add_paragraph("Para 3 content.")
        doc.save(str(doc_path))

        merger = DocumentMerger(output_dir=temp_dir)
        corrections = [
            {
                "original": "Para",
                "fixed": "Modified",
                "paragraph_index": 2,
                "context_before": "1",
            }
        ]

        result = merger.merge(doc_path, corrections)
        doc = Document(str(result))
        texts = [p.text for p in doc.paragraphs]

        # Only paragraph 2 should be modified (contains "Para 2")
        assert texts[0] == "Para 1 content."
        assert texts[1] == "Modified 2 content."
        assert texts[2] == "Para 3 content."

    def test_context_before_used_for_disambiguation(self, temp_dir):
        """Test context_before is used to find correct location."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not available")

        doc_path = temp_dir / "context_test.docx"
        doc = Document()
        doc.add_paragraph("unique marker first occurrence.")
        doc.add_paragraph("some other text unique marker here.")
        doc.save(str(doc_path))

        merger = DocumentMerger(output_dir=temp_dir)
        corrections = [
            {
                "original": "unique marker",
                "fixed": "REPLACED",
                "paragraph_index": 1,
                "context_before": "first",
            }
        ]

        result = merger.merge(doc_path, corrections)
        doc = Document(str(result))
        texts = [p.text for p in doc.paragraphs]

        # Only first paragraph should be modified
        assert texts[0] == "REPLACED first occurrence."
        assert "REPLACED" not in texts[1]

    def test_context_after_used_for_disambiguation(self, temp_dir):
        """Test context_after is used to find correct location."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not available")

        doc_path = temp_dir / "context_after_test.docx"
        doc = Document()
        doc.add_paragraph("text before unique word after.")
        doc.add_paragraph("text before unique word elsewhere.")
        doc.save(str(doc_path))

        merger = DocumentMerger(output_dir=temp_dir)
        corrections = [
            {
                "original": "unique word",
                "fixed": "CHANGED",
                "paragraph_index": 1,
                "context_after": "after",
            }
        ]

        result = merger.merge(doc_path, corrections)
        doc = Document(str(result))
        texts = [p.text for p in doc.paragraphs]

        # Only first paragraph should be modified
        assert texts[0] == "text before CHANGED after."
        assert "CHANGED" not in texts[1]

    def test_ai_word_skill_detection(self):
        """Test AI-Word-Skill availability detection."""
        merger = DocumentMerger()
        # Should be False since we haven't mocked the import
        assert merger.using_ai_word_skill is False

    def test_fuzzy_match_threshold(self):
        """Test fuzzy matching threshold behavior."""
        from app.infrastructure.docx.document_merger import DocumentMerger

        merger = DocumentMerger()

        # Exact match
        assert merger._fuzzy_match("hello", "hello world", 0.8) is True

        # Close match
        assert merger._fuzzy_match("hello", "hello", 0.8) is True

        # No match
        assert merger._fuzzy_match("xyz", "abc", 0.8) is False

        # Empty patterns
        assert merger._fuzzy_match("", "abc", 0.8) is False
        assert merger._fuzzy_match("abc", "", 0.8) is False


class TestDocumentMergerAIWordSkillFallback:
    """Test fallback behavior when AI-Word-Skill is unavailable."""

    @pytest.fixture
    def temp_dir(self):
        """Create a temporary directory for test files."""
        temp_path = Path(tempfile.mkdtemp())
        yield temp_path
        shutil.rmtree(temp_path, ignore_errors=True)

    @pytest.fixture
    def sample_docx(self, temp_dir):
        """Create a sample .docx file for testing."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not available")

        doc_path = temp_dir / "sample.docx"
        doc = Document()
        doc.add_paragraph("Test paragraph content.")
        doc.save(str(doc_path))
        return doc_path

    def test_fallback_to_basic_replacement(self, temp_dir, sample_docx):
        """Test that fallback to basic replacement works."""
        merger = DocumentMerger(output_dir=temp_dir)

        # Even if _ai_word_skill_available is True, if the import fails,
        # it should fall back to basic replacement
        corrections = [
            {"original": "Test", "fixed": "Replaced", "paragraph_index": 1}
        ]

        result = merger.merge(sample_docx, corrections)
        assert result.exists()

        from docx import Document
        doc = Document(str(result))
        assert "Replaced paragraph content." in [p.text for p in doc.paragraphs]


if __name__ == "__main__":
    pytest.main([__file__, "-v"])